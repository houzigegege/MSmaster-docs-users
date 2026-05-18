import argparse
import os
import re
import shutil
from dataclasses import dataclass
from typing import Dict, Iterator, List, Optional, Tuple, Union

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


Block = Union[Paragraph, Table]


def iter_block_items(doc: Document) -> Iterator[Block]:
    """
    按文档顺序迭代 block：Paragraph 和 Table。
    python-docx 默认 doc.paragraphs / doc.tables 会打乱顺序，所以这里需要自己迭代 body。
    """
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif tag.endswith("}tbl"):
            yield Table(child, doc)


def extract_paragraph_text_with_images(
    paragraph: Paragraph,
    img_dir: str,
    rel_image_cache: Dict[str, str],
    img_counter: int,
) -> Tuple[str, int]:
    """
    将段落内容转成 markdown 可用文本，并把段落内出现的图片插入到文本中。
    """
    parts: List[str] = []

    # 使用 run 级别遍历，能较好保持段落内的图片位置。
    for run in paragraph.runs:
        if run.text:
            parts.append(run.text)

        # 查找该 run 内的图片 blip
        blips = run._element.xpath(".//a:blip")
        for blip in blips:
            r_id = blip.get(qn("r:embed"))
            if not r_id:
                continue

            if r_id in rel_image_cache:
                rel_path = rel_image_cache[r_id]
                parts.append(f"![]({rel_path})")
                continue

            related = paragraph.part.related_parts
            if r_id not in related:
                continue

            image_part = related[r_id]
            content_type = getattr(image_part, "content_type", "")

            # 尽量从 content_type 推断扩展名
            ext = "png"
            if content_type.endswith("jpeg"):
                ext = "jpg"
            elif content_type.endswith("png"):
                ext = "png"
            elif content_type.endswith("gif"):
                ext = "gif"
            elif content_type.endswith("bmp"):
                ext = "bmp"
            elif content_type.endswith("tiff"):
                ext = "tif"

            img_counter += 1
            filename = f"img_{img_counter:04d}.{ext}"
            abs_path = os.path.join(img_dir, filename)
            os.makedirs(img_dir, exist_ok=True)

            with open(abs_path, "wb") as f:
                f.write(image_part.blob)

            # mkdocs 对 docs 下的路径是可用的，统一用相对路径（正斜杠）
            rel_path = os.path.join("assets", "manual", filename).replace("\\", "/")
            rel_image_cache[r_id] = rel_path
            parts.append(f"![]({rel_path})")

    text = "".join(parts).strip()
    return text, img_counter


def table_to_markdown(
    table: Table,
    img_dir: str,
    rel_image_cache: Dict[str, str],
    img_counter: int,
) -> Tuple[str, int]:
    """
    将 docx 表格转成 markdown。
    注意：docx 表格在复杂合并单元格场景下可能无法 100% 还原，这里做的是可用的近似转换。
    """
    rows: List[List[str]] = []
    for row in table.rows:
        cells: List[str] = []
        for cell in row.cells:
            # 尽量保留 cell 内的图片：不直接依赖 cell.text（可能会丢掉图片）。
            cell_parts: List[str] = []
            for p in cell.paragraphs:
                p_text, img_counter = extract_paragraph_text_with_images(
                    paragraph=p,
                    img_dir=img_dir,
                    rel_image_cache=rel_image_cache,
                    img_counter=img_counter,
                )
                p_text = p_text.strip()
                if p_text:
                    cell_parts.append(p_text)

            cell_md = "<br>".join(cell_parts).strip()
            cells.append(cell_md)
        rows.append(cells)

    # 估计列数：使用第一行列数
    if not rows or not rows[0]:
        return ""

    col_count = len(rows[0])
    # 防御：如果后面行列数不同，截断/补空
    def normalize_row(r: List[str]) -> List[str]:
        r2 = r[:col_count]
        if len(r2) < col_count:
            r2 += [""] * (col_count - len(r2))
        return r2

    header = normalize_row(rows[0])
    body = [normalize_row(r) for r in rows[1:]]

    def md_row(cells: List[str]) -> str:
        return "| " + " | ".join(cells) + " |"

    sep = "| " + " | ".join(["---"] * col_count) + " |"
    md_lines = [md_row(header), sep]
    md_lines.extend(md_row(r) for r in body)
    return "\n".join(md_lines), img_counter


@dataclass
class Page:
    title: str
    filename: str
    lines: List[str]


def strip_section_number_prefix(title: str) -> str:
    """
    Remove leading section numbering from Word headings, e.g.
    - '4. Quick Start', '5.Fusion ...'  -> drop 'N.'
    - '4.1 I Have ...', '3.2.1 Foo'     -> drop 'N.M...' (subsection indices)
    - '1) Overview'                     -> drop 'N)'
    """
    t = (title or "").strip()
    if not t:
        return t
    # Multi-part indices first (4.1, 4.2, 3.2.1) so we do not split "4.1" as "4." + "1"
    t = re.sub(r"^\d+(?:\.\d+)+\s+", "", t)
    # Single index + dot: "4. " or "12." (also matches "5.Fusion" via \s* = zero spaces)
    t = re.sub(r"^\d+\.\s*", "", t)
    # "1) " style
    t = re.sub(r"^\d+\)\s*", "", t)
    return t.strip()


def safe_first_title_from_doc(doc: Document) -> str:
    """
    尝试从正文找文档标题：优先 Heading 1；否则用第一段非空 Normal。
    """
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        s = (p.style.name or "").strip() if p.style else ""
        if s.lower().startswith("heading"):
            return t
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        return t
    return "Manual"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="Manual.docx", help="Manual.docx 路径（默认：当前目录 Manual.docx）")
    parser.add_argument(
        "--manual-dir",
        default="docs/manual",
        help="要输出到 mkdocs 的“自动生成手册”目录（默认：docs/manual）",
    )
    parser.add_argument(
        "--clean-manual",
        action="store_true",
        help="是否清空 manual-dir 后重建（仅影响 docs/manual）",
    )
    parser.add_argument(
        "--write-mkdocs",
        action="store_true",
        help="是否同步生成/覆盖 mkdocs.yml（不推荐；方案 A 下默认不写）",
    )
    parser.add_argument("--output-mkdocs", default="mkdocs.yml", help="输出 mkdocs.yml 路径（默认：mkdocs.yml）")
    args = parser.parse_args()

    input_path = args.input
    output_mkdocs = args.output_mkdocs
    manual_dir = args.manual_dir

    if not os.path.exists(input_path):
        raise FileNotFoundError(f"找不到输入文件：{input_path}")

    if args.clean_manual and os.path.exists(manual_dir):
        shutil.rmtree(manual_dir)

    os.makedirs(manual_dir, exist_ok=True)

    # 图片目录（输出到 docs/manual/assets/manual，保证 markdown 的相对引用可用）
    img_dir = os.path.join(manual_dir, "assets", "manual")
    os.makedirs(img_dir, exist_ok=True)

    doc = Document(input_path)

    title = safe_first_title_from_doc(doc)

    pages: List[Page] = []
    current: Optional[Page] = None

    # 用于图片去重与编号连续
    rel_image_cache: Dict[str, str] = {}
    img_counter = 0

    heading_re = re.compile(r"^Heading\s+(\d+)\s*$", re.IGNORECASE)

    def start_new_page(h1_title: str, idx: int) -> Page:
        clean = strip_section_number_prefix(h1_title)
        filename = f"section_{idx:02d}.md"
        p = Page(title=clean, filename=filename, lines=[f"# {clean}", ""])
        return p

    # Keep in sync with GitHub Releases (update URL when publishing a new release).
    _download_url = (
        "https://github.com/houzigegege/MSmaster-docs-users/releases/download/"
        "v1.0.0/MSmaster_V1.0.0.7z"
    )
    install_download_block = "\n".join(
        [
            "MSmaster for **Windows 10 / 11 (64-bit)**. You need about **1.0 GB** download "
            "space and **4.3 GB** free disk after extraction. **8 GB RAM** minimum; "
            "**16 GB** recommended.",
            "",
            "## Download",
            "",
            f"[MSmaster_V1.0.0.7z]({_download_url})",
            "",
            "## Installation",
            "",
            "1. Download `MSmaster_V1.0.0.7z` using the link above.",
            "2. Extract the archive to a local folder (for example `C:\\MSmaster\\`).",
            "3. Open the extracted folder and run **`MSmaster.exe`** (or the main launcher "
            "shown in the folder).",
            "4. On first launch, allow Windows to run the application if a security prompt "
            "appears (only if you trust this release source).",
            "5. For workflows, parameters, and result interpretation, continue with the "
            "[Scientific Usage Guide](https://houzigegege.github.io/MSmaster-docs-users/) — "
            "start from [Quick Start](https://houzigegege.github.io/MSmaster-docs-users/manual/section_04/) "
            "if you are new to the platform.",
            "",
            "!!! note",
            "    Identification and networking outputs are computational predictions. "
            "Validate important findings with standards, confirmatory MS/MS, and appropriate "
            "experimental controls.",
            "",
        ]
    )

    about_license_affiliations_block = "\n".join(
        [
            "## License: MIT",
            "",
            "This software is released under the MIT License.",
            "",
            "```text",
            "Copyright (c) 2026 MSmaster Contributors",
            "",
            "Permission is hereby granted, free of charge, to any person obtaining a copy",
            "of this software and associated documentation files (the \"Software\"), to deal",
            "in the Software without restriction, including without limitation the rights",
            "to use, copy, modify, merge, publish, distribute, sublicense, and/or sell",
            "copies of the Software, and to permit persons to whom the Software is",
            "furnished to do so, subject to the following conditions:",
            "",
            "The above copyright notice and this permission notice shall be included in all",
            "copies or substantial portions of the Software.",
            "",
            "THE SOFTWARE IS PROVIDED \"AS IS\", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR",
            "IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,",
            "FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE",
            "AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER",
            "LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,",
            "OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE",
            "SOFTWARE.",
            "```",
            "",
            "## Departmental affiliations",
            "",
            "a Department of Biochemistry, Max Planck Institute for Chemical Ecology, Hans-Knöll-Straße 8, 07745 Jena, Germany",
            "b State Key Laboratory of Phytochemistry and Plant Resources in West China, Kunming Institute of Botany, Chinese Academy of Sciences, Kunming 650201, Yunnan, China",
            "",
        ]
    )

    def _strip_install_download_block(lines: List[str]) -> List[str]:
        """Remove injected download/install text before re-injecting; keep screenshot onward."""
        if not lines:
            return lines
        out = [lines[0]]
        for i in range(1, len(lines)):
            if lines[i].strip().startswith("![]"):
                return out + lines[i:]
        return out

    def maybe_inject_install_download(page: Page) -> None:
        """
        Ensure the Install section always contains the installer download link.
        This prevents manual edits from being overwritten when regenerating from Manual.docx.
        """
        is_install_section = "install" in (page.title or "").lower()
        if not is_install_section:
            return
        page.lines = _strip_install_download_block(page.lines)
        insert_at = 1 if len(page.lines) > 1 else len(page.lines)
        page.lines[insert_at:insert_at] = [install_download_block, ""]

    def maybe_inject_about_license(page: Page) -> None:
        """
        Ensure the About section always contains license + affiliation info.
        """
        if page.filename != "section_01.md":
            return
        md = "\n".join(page.lines)
        if "License: MIT" in md and "Departmental affiliations" in md:
            return
        insert_at = 1 if len(page.lines) > 1 else len(page.lines)
        page.lines[insert_at:insert_at] = [about_license_affiliations_block, ""]

    h1_count = 0

    # 迭代文档 block，尽量保持表格/图片顺序
    for block in iter_block_items(doc):
        if isinstance(block, Table):
            md, img_counter = table_to_markdown(
                block,
                img_dir=img_dir,
                rel_image_cache=rel_image_cache,
                img_counter=img_counter,
            )
            if not md.strip():
                continue
            if current is None:
                # 表格出现在 Heading 1 之前时，先丢到首页
                if not pages:
                    pages.append(Page(title="Overview", filename="index.md", lines=[f"# {title}", ""]))
                    current = pages[0]
                else:
                    current = pages[0]
            current.lines.append(md)
            current.lines.append("")
            continue

        # Paragraph
        style_name = (block.style.name or "").strip() if block.style else ""
        raw_text, img_counter = extract_paragraph_text_with_images(
            paragraph=block,
            img_dir=img_dir,
            rel_image_cache=rel_image_cache,
            img_counter=img_counter,
        )

        text = raw_text.strip()
        if not text:
            continue

        m = heading_re.match(style_name)
        if m:
            level = int(m.group(1))
            if level == 1:
                h1_count += 1
                current = start_new_page(text, h1_count)
                pages.append(current)
            else:
                if current is None:
                    # Heading 2/3/4 出现在 Heading 1 之前：放到临时首页
                    if not pages:
                        pages.append(Page(title="Overview", filename="index.md", lines=[f"# {title}", ""]))
                        current = pages[0]
                    else:
                        current = pages[0]
                hashes = "#" * level
                clean_heading = strip_section_number_prefix(text)
                current.lines.append(f"{hashes} {clean_heading}")
                current.lines.append("")
        else:
            if current is None:
                # 正文出现在 Heading 1 之前：放到首页
                if not pages:
                    pages.append(Page(title="Overview", filename="index.md", lines=[f"# {title}", ""]))
                    current = pages[0]
                else:
                    current = pages[0]
            current.lines.append(text)
            current.lines.append("")

    # 生成 index.md（放到自动生成手册目录内）
    index_path = os.path.join(manual_dir, "index.md")
    # 只为 Heading 1 页面生成目录；忽略临时 Overview 内容（可以自行改造）
    section_pages = [p for p in pages if p.filename.startswith("section_") or p.filename == "index.md"]
    section_pages = [p for p in section_pages if p.filename != "index.md"]

    index_lines: List[str] = [
        f"# {title}",
        "",
        "## Contents",
        "",
    ]
    for p in section_pages:
        # mkdocs 的 nav 里也会展示，这里做一个简洁目录
        index_lines.append(f"- [{p.title}]({p.filename})")
    index_lines.append("")
    # 如果没有 Heading 1，那么就把 Overview 当首页
    overview_candidates = [p for p in pages if p.filename == "index.md"]
    if not section_pages and overview_candidates:
        index_lines = overview_candidates[0].lines

    with open(index_path, "w", encoding="utf-8") as f:
        f.write("\n".join(index_lines))

    # 写入 section pages
    for p in pages:
        if p.filename == "index.md":
            continue
        maybe_inject_install_download(p)
        maybe_inject_about_license(p)
        out_path = os.path.join(manual_dir, p.filename)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(p.lines).strip() + "\n")

    # 方案 A：mkdocs.yml 由你手工设计/维护；自动生成脚本默认不覆盖
    if args.write_mkdocs:
        nav_entries: List[str] = []
        nav_entries.append(f'  - 主页: manual/index.md')
        for p in [p for p in pages if p.filename.startswith("section_")]:
            nav_entries.append(f"  - {p.title}: manual/{p.filename}")

        mkdocs_yml = "\n".join(
            [
                f"site_name: {title}",
                "theme:",
                "  name: material",
                "  language: zh",
                "  features:",
                "    - navigation.tabs",
                "    - navigation.sections",
                "    - navigation.top",
                "    - search.highlight",
                "plugins:",
                "  - search",
                "markdown_extensions:",
                "  - pymdownx.highlight",
                "  - tables",
                "nav:",
                *nav_entries,
                "",
            ]
        )
        with open(output_mkdocs, "w", encoding="utf-8") as f:
            f.write(mkdocs_yml)
        print(f"[OK] 已生成：{output_mkdocs}")

    gen_count = len([p for p in pages if p.filename.startswith("section_")])
    print(f"[OK] 已生成自动手册：{manual_dir}/index.md 和 {gen_count} 个分章节页面")


if __name__ == "__main__":
    main()

